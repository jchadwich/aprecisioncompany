from docx import Document
from docx.shared import Pt
from docx.shared import Inches

# from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import folium
import pandas as pd
from math import ceil
from math import floor
from pandas import *
from openpyxl import load_workbook
from datetime import datetime
import geopy
from folium.plugins import FloatImage
from geopy.geocoders import Nominatim, ArcGIS, GoogleV3  # Geocoder APIs
import psycopg2

# data from db
conn = psycopg2.connect(
    database="aprecisioncompanydb",
    user="postgres",
    password="infrastructure",
    host="127.0.0.1",
    port="5432",
)
conn.autocommit = True
cursor = conn.cursor()

# Create sorted list of id values to get most recent id for cursor.execute
cursor.execute("""SELECT id from docgen_docgen""")
i = cursor.fetchall()
n = 0  # Index variable
idList = []
for item in i:
    # Pull int values from lists
    idList.append(i[n][0])
    n += 1
print(idList[-1])
cursor.execute(
    f"""SELECT * from docgen_docgen WHERE id = {idList[-1]}"""
)  # Most recent entry only
result = cursor.fetchone()
csvFile = result[7]
print(pd.read_csv(csvFile))


# Ensure that the input file is a csv
try:
    userInFile = result[7]
except:
    userInFile = "Delta_Ridge_Townhomes_139926-21Apr2022.csv"

oidList = userInFile["No."].values.tolist()
lenList = userInFile["Length"].values.tolist()
widList = userInFile["Width"].values.tolist()
scList = userInFile["Special Case"].values.tolist()
qdList = userInFile["Quick Description"].values.tolist()
xList = userInFile["x"].values.tolist()
yList = userInFile["y"].values.tolist()

try:
    saveFile = ""  # result[]
except:
    saveFile = "Delta_Ridge_Townhomes_139926-21Apr2022_proposal.docx"
try:
    inputEntity = result[1]
except:
    inputEntity = "Delta Ridge West and South"
try:
    PPRNum = result[26]
except:
    PPRNum = "11422-211-01"
try:
    PONum = result[25]
except:
    PONum = ""
try:
    indivName = result[10]
except:
    indivName = "Bonnie Giles"
try:
    indivTitle = result[12]
except:
    indivTitle = "Association Manager"
try:
    BDName = result[14]
except:
    BDName = "Jack"
try:
    BDTitle = result[16]
except:
    BDTitle = "BD"
try:
    BDExt = result[13]
except:
    BDExt = "300"
try:
    BDPh = result[15]
except:
    BDPh = "(111)111-000"
try:
    contactAddress = result[8]
except:
    contactAddress = "8305 Falls of Neuse Road, Suite 200, Raleigh, NC 27615"
try:
    contactPh = result[11]
except:
    contactPh = "919-878-8787 ext. 255"
try:
    contactEmail = result[9]
except:
    contactEmail = "bonitagiles@towneproperties.com"
try:
    projName = result[27]
except:
    projName = "Wilmington Proj"
try:
    specs = result[32]
except:
    specs = "½” to 2½”"
try:
    segwaysUsed = result[30]
except:
    segwaysUsed = "no"
try:
    dnrPrice = result[18]
except:
    dnrPrice = 15
try:
    city = result[2]
except:
    city = inputEntity
try:
    safetyIncident = result[29]
except:
    safetyIncident = "0"
try:
    smPrice = result[31]
except:
    smPrice = 2.5
try:
    mdPrice = result[23]
except:
    mdPrice = 5
try:
    lgPrice = result[20]
except:
    lgPrice = 10
try:
    curbPrice = result[17]
except:
    curbPrice = 10
try:
    expNumDaysToRepair = ""  # result[]
except:
    expNumDaysToRepair = 3
try:
    expNumTechs = result[33]
except:
    expNumTechs = 3
try:
    pssMin = result[28]
except:
    pssMin = "$5,000"
try:
    avgPrice = ""  # result[]
except:
    avgPrice = 2.50
try:
    dtcLow = result[24]
except:
    dtcLow = 1
try:
    dtcHigh = result[22]
except:
    dtcHigh = 2
try:
    BDEmail = ""  # result[]
except:
    BDEmail = "bd@precisionsafesidewalks.com"
try:
    workOrderLoc = result[34]
except:
    workOrderLoc = "No"
try:
    mapLayers = result[21]
except:  # ACCEPTABLE INPUT: "dnr", "repairs", "dnr and repair"
    mapLayers = "dnr and repair"


inFile = "C:/aprecisioncompany/" + userInFile

doc = Document("PPR TEMPLATE 11-11-2022 TM.docx")
section = doc.sections[0]
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = paragraph.text.replace("Entity", inputEntity)

doc_header = True

centerX = 0
centerY = 0
# Calculate center of map from averages of lats and longs
xList = []
yList = []

# Reverse Geocode instantiation
g = ArcGIS()
addresses = []

x = pd.read_excel(userInFile, sheet_name="1", usecols="V")
y = pd.read_excel(userInFile, sheet_name="1", usecols="W")
xl = x.values.tolist()
yl = y.values.tolist()
xList = []
yList = []
i = 0
for item in xl:
    if i >= 24:
        xList.append(item[0])
    i += 1
i = 0
for item in yl:
    if i >= 24:
        yList.append(item[0])
    i += 1

print(xList)
print(yList)
centerX = sum(xList) / (len(xList))
centerY = sum(yList) / (len(yList))
print(centerX, centerY)

# **********************************************************************************************************************
count = pd.read_excel(userInFile, sheet_name="SUMMARY", usecols="F")
countL = count.values.tolist()
smCount = countL[5]
smCount = int(smCount[0])
mdCount = countL[6]
mdCount = int(mdCount[0])
lgCount = countL[7]
lgCount = int(lgCount[0])
#
repSqft = pd.read_excel(userInFile, sheet_name="SUMMARY", usecols="S")
sqftR = repSqft.values.tolist()
dnrSqft = sqftR[5]
dnrSqft = int(dnrSqft[0])
#
curbLF = pd.read_excel(userInFile, sheet_name="SUMMARY", usecols="O")
lfCurb = curbLF.values.tolist()
curbLenFT = lfCurb[5]
curbLenFT = int(curbLenFT[0])
#
curbC = pd.read_excel(userInFile, sheet_name="SUMMARY", usecols="P")
cCurb = curbC.values.tolist()
curbCost = cCurb[5]
curbCost = int(curbCost[0])
#
price_opt1 = pd.read_excel(userInFile, sheet_name="SUMMARY", usecols="H")
pOpt1 = price_opt1.values.tolist()
p1Sm = pOpt1[5]
p1Sm = int(p1Sm[0])
p1Md = pOpt1[6]
p1Md = int(p1Md[0])
p1Lg = pOpt1[7]
p1Lg = int(p1Lg[0])
#
price_opt2 = pd.read_excel(userInFile, sheet_name="SUMMARY", usecols="G")
pOpt2 = price_opt2.values.tolist()
p2Sm = pOpt2[5]
p2Sm = int(p2Sm[0])
p2Md = pOpt2[6]
p2Md = int(p2Md[0])
p2Lg = pOpt2[7]
p2Lg = int(p2Lg[0])
#
sSqft = pd.read_excel(userInFile, sheet_name="1", usecols="AE")
sqftSm = sSqft.values.tolist()
smSqft = sqftSm[8]
smSqft = int(smSqft[0])
mSqft = pd.read_excel(userInFile, sheet_name="1", usecols="AF")
sqftMd = mSqft.values.tolist()
mdSqft = sqftMd[8]
mdSqft = int(mdSqft[0])
lSqft = pd.read_excel(userInFile, sheet_name="1", usecols="AG")
sqftLg = lSqft.values.tolist()
lgSqft = sqftLg[8]
lgSqft = int(lgSqft[0])
# **********************************************************************************************************************


def Doc():
    # dnrCount = 0
    # curbCount = 0
    totalCount = 0
    tripHazardCount = smCount + mdCount + lgCount

    totalSqft = dnrSqft + smSqft + mdSqft + lgSqft  # Does not include curb
    curbLength = curbLenFT
    tripHazardSqft = lgSqft + mdSqft + smSqft

    lgCost = p1Lg  # based on price option 1
    mdCost = p1Md  # based on price option 1
    smCost = p1Sm  # based on price option 1
    dnrCost = dnrSqft * dnrPrice
    totalCost = 0
    totalDnrCost = 0  # Total cost if all deficiencies were treated as dnr

    wasteConcreteLow = (
        0  # How many tons were saved from going to the landfill? Low number
    )
    wasteConcreteHigh = (
        0  # How many tons were saved from going to the landfill? High number
    )
    wastConcreteCF = 0  # How many cubic feet of natural resources were saved?
    gasSaved = 0  # How many gallons of gasoline were saved?
    tonsCO2 = 0  # How many tons of CO2 were saved from being pumped into atmosphere?
    # *********************************************************************************************************************

    # **********************************************************************************************************************
    totalCost = dnrCost + smCost + mdCost + lgCost + curbCost
    totalHazardCost = smCost + mdCost + lgCost + curbCost
    print("Total Hazard Cost: " + str(totalHazardCost))

    # Calculate cost if dnr only was used
    print(lgSqft)
    print(mdSqft)
    print(smSqft)
    print(dnrSqft)
    print(curbLength)
    totalDnrCost = ((lgSqft + mdSqft + smSqft + dnrSqft) * dnrPrice) + (45 * curbLength)
    print(totalDnrCost)
    s = pd.read_excel(userInFile, sheet_name="GREEN SAVINGS", usecols="E")
    sList = s.values.tolist()
    savings = sList[54]
    savings = int(savings[0])
    print("savings are:")
    print(savings)

    greenSavings = pd.read_excel(userInFile, sheet_name="GREEN SAVINGS", usecols="C")
    greenSavingsList = greenSavings.values.tolist()

    # Calculate Natural Resource Savings
    tcf = greenSavingsList[6]  # total cubic feet
    tcf = float(tcf[0])
    tl = greenSavingsList[12]  # tons (low)
    tl = float(tl[0])
    th = greenSavingsList[13]  # tons (high)
    th = float(th[0])
    dump = ceil(th) * 8  # total miles to dump dnr
    bh = ceil(th) * 8  # backhoe transport miles
    totMiles = dump + bh + 20  # total miles
    ga = greenSavingsList[29]  # gallons required to transport to new site
    ga = float(ga[0])
    gan = totMiles / 8  # gallons required to transport to new site
    miscMi = (ceil(dump) / 4) * 6  # miscellaneous mileage
    wf = miscMi / 15  # fuel for workers
    totFuelSaved = wf + ga + gan  # total gallons fossil fuels avoided

    cpCO2 = ((th + tl) / 2) * 0.13078  # CO2 avoided during concrete production
    ffCO2 = totFuelSaved * 0.009  # CO2 avoided fossil fuels
    totCO2 = greenSavingsList[34]  # total CO2 avoided
    totCO2 = float(totCO2[0])
    # *********************************************************************************************************************

    for paragraph in doc.paragraphs:
        if "ENTITY" in paragraph.text:
            paragraph.text = paragraph.text.replace("ENTITY", inputEntity)
        elif "the Entity" in paragraph.text:
            paragraph.text = paragraph.text.replace("the Entity", inputEntity)
        if "Name, Title" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "Name, Title", indivName + ", " + indivTitle
            )
        if "BDM, Title" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "BDM, Title", BDName + ", " + BDTitle
            )
        if "305" in paragraph.text:
            paragraph.text = paragraph.text.replace("305", BDExt)
        if "(xxx) xxx-xxxx" in paragraph.text:
            paragraph.text = paragraph.text.replace("(xxx) xxx-xxxx", BDPh)
        if "Address" in paragraph.text:
            paragraph.text = paragraph.text.replace("Address", contactAddress)
        if "Contact Phone | Contact email" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "Contact Phone | Contact email", contactPh + " | " + contactEmail
            )
        if "the project name" in paragraph.text:
            paragraph.text = paragraph.text.replace("the project name", projName)
        if "½” to 2½”" in paragraph.text:
            paragraph.text = paragraph.text.replace("½” to 2½” ", specs)
        # if segwaysUsed == "no":
        # if "While a portion of the area was covered using Segways, some of the areas had to be covered on foot because of the hazard density and weather." in paragraph.text:
        # paragraph.text = paragraph.text.replace("While a portion of the area was covered using Segways, some of the areas had to be covered on foot because of the hazard density and weather.", "All areas had to be covered on foot because of the hazard density and weather")
        if "15.00/sqft " in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "15.00/sqft", str(dnrPrice) + ".00/sqft"
            )  # not updating - inside text box
        if "15.00 per" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "15.00 per", str(dnrPrice) + ".00 per"
            )
        if "City" in paragraph.text:
            paragraph.text = paragraph.text.replace("City", city)
        if "There were no safety incidents" in paragraph.text:
            if safetyIncident == 1:
                paragraph.text = paragraph.text.replace(
                    "There Were no safety incidents", "There was 1 safety incident"
                )
            elif safetyIncident == 0:
                paragraph.text = paragraph.text.replace(
                    "There Were no safety incidents", "There Were no safety incidents"
                )
            else:
                paragraph.text = paragraph.text.replace(
                    "There were no safety incidents",
                    "There were " + str(safetyIncident) + " safety incidents",
                )
        if "340" in paragraph.text:
            paragraph.text = paragraph.text.replace("340", str(int(tripHazardCount)))
        if "tripHazardSqft" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "tripHazardSqft", str(int(tripHazardSqft))
            )
        if "curbLength" in paragraph.text:
            paragraph.text = paragraph.text.replace("curbLength", str(int(curbLength)))
        if "curbCost" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "curbCost", str(curbCost)
            )  # Evenually check if larger than 1,000, if so, add comma
        if "savings_" in paragraph.text:
            paragraph.text = paragraph.text.replace("savings_", str(savings))
        if "totalDnrCost" in paragraph.text:
            paragraph.text = paragraph.text.replace("totalDnrCost", str(totalDnrCost))
        if "XXX22-211-01" in paragraph.text:
            paragraph.text = paragraph.text.replace("XXX22-211-01", PPRNum)
        if "totalHazardCost" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "totalHazardCost", str(totalHazardCost)
            )
        if "totCO2" in paragraph.text:
            paragraph.text = paragraph.text.replace("totCO2", str(totCO2))
        if "TL" in paragraph.text:
            paragraph.text = paragraph.text.replace("TL", str(tl))
        if "TH" in paragraph.text:
            paragraph.text = paragraph.text.replace("TH", str(th))
        if "totFuelSaved" in paragraph.text:
            paragraph.text = paragraph.text.replace("totFuelSaved", str(totFuelSaved))
        if "_cos_" in paragraph.text:
            paragraph.text = paragraph.text.replace("_cos_", str(totalHazardCost))
        if "_sav_" in paragraph.text:
            paragraph.text = paragraph.text.replace("_sav_", str(savings))
        if "Change $15.00 above to dnrPricing" in paragraph.text:
            if dnrPrice == 15:
                paragraph.text = paragraph.text.replace(
                    "Change $15.00 above to dnrPricing", ""
                )
            else:
                paragraph.text = paragraph.text.replace(
                    "Change $15.00 above to dnrPricing",
                    "Change $15.00 above to $" + str(dnrPrice) + ".00",
                )

    tables = doc.tables
    # Update table 1
    # Location
    table = doc.tables[0]
    table.cell(2, 0).text = str(inputEntity)
    # Hazard Counts
    table = doc.tables[0]
    table.cell(2, 2).text = str(smCount)
    table = doc.tables[0]
    table.cell(3, 2).text = str(mdCount)
    table = doc.tables[0]
    table.cell(4, 2).text = str(lgCount)
    table = doc.tables[0]
    table.cell(5, 2).text = str(smCount + mdCount + lgCount)
    # Hazard Costs
    table = doc.tables[0]
    table.cell(2, 3).text = "$" + str(smCost)
    table = doc.tables[0]
    table.cell(3, 3).text = "$" + str(mdCost)
    table = doc.tables[0]
    table.cell(4, 3).text = "$" + str(lgCost)
    table = doc.tables[0]
    table.cell(5, 3).text = "$" + str(smCost + mdCost + lgCost)
    # Update table 2
    # Location
    table = doc.tables[1]
    table.cell(2, 0).text = str(inputEntity)
    # Curb lengths
    table = doc.tables[1]
    table.cell(2, 1).text = str(int(curbLength))
    table = doc.tables[1]
    table.cell(3, 1).text = str(int(curbLength))
    # Curb Costs
    table = doc.tables[1]
    table.cell(2, 2).text = "$" + str(curbCost)
    table = doc.tables[1]
    table.cell(3, 2).text = "$" + str(curbCost)

    # Update ppr calculations table (page 8)
    # table = doc.tables[2]
    # table.cell(2,1).text = "Based on demolition and replacement (D&R) costs of $" + str(dnrPrice) +".00 including:"

    # Add numbers for natural resource savings to table
    table = doc.tables[3]
    table.cell(9, 0).text = inputEntity
    table = doc.tables[3]
    table.cell(9, 1).text = str(ceil(tcf))
    table = doc.tables[3]
    table.cell(9, 2).text = str(ceil(tl))
    table = doc.tables[3]
    table.cell(9, 3).text = str(ceil(th))
    table = doc.tables[3]
    table.cell(9, 4).text = str(ceil(ga))
    table = doc.tables[3]
    table.cell(9, 5).text = str(round(totCO2, 1))

    # Add pictures
    p = tables[4].rows[0].cells[0].add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture("sw.jpg")

    p = tables[5].rows[0].cells[0].add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture("sw2.jpg")

    # Add test pictures - 4 on 1 page
    p = tables[6].rows[0].cells[0].add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture("test.png", width=Inches(2.625), height=Inches(3.5))
    p = tables[6].rows[0].cells[1].add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture("test.png", width=Inches(2.625), height=Inches(3.5))
    p = tables[6].rows[1].cells[0].add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture("test.png", width=Inches(2.625), height=Inches(3.5))
    p = tables[6].rows[1].cells[1].add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture("test.png", width=Inches(2.625), height=Inches(3.5))

    print("dnr cost if all were treated as dnr: " + "$" + str(totalDnrCost))
    print("trip hazard cost: " + "$" + str(totalCost))

    print(totalSqft)
    doc.save(saveFile)


d = Doc()
